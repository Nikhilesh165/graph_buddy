"""Regenerates the two binary sample sources (company_handbook.docx and
quarterly_report.pdf) in this directory. The other three sample sources
(employees.csv, meeting_notes.txt, project_overview.md) are plain text and
committed directly -- no generator needed.

Run from the repo root with the backend's environment (python-docx and
fpdf2 are already backend dependencies):

    cd backend && uv run python ../examples/sample-sources/generate_binary_samples.py
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt
from fpdf import FPDF
from fpdf.enums import XPos, YPos

OUT_DIR = Path(__file__).resolve().parent


def build_docx(out_path: Path) -> None:
    doc = Document()

    doc.add_heading("Solstice Robotics -- Employee Handbook", level=0)
    doc.add_paragraph(
        "Solstice Robotics, Inc. is headquartered in Austin, Texas, with a "
        "field office in Des Moines, Iowa. This handbook applies to all "
        "full-time employees."
    )

    doc.add_heading("1. Company Leadership", level=1)
    doc.add_paragraph(
        "Maria Chen (CEO & Co-Founder) and David Okoye (CTO & Co-Founder) "
        "founded Solstice Robotics in 2021. Maria oversees Sales and "
        "Operations; David oversees Engineering."
    )
    doc.add_paragraph(
        "Priya Nair, VP of Engineering, reports to David Okoye and leads "
        "the Engineering department, including Project Meridian and "
        "Project Halcyon. Tom Alvarez, Head of Sales, reports to Maria "
        "Chen. Grace Kim, People Operations Lead, reports to Maria Chen."
    )

    doc.add_heading("2. Locations", level=1)
    doc.add_paragraph(
        "Austin HQ houses Executive, Sales, and most of Operations. The "
        "Des Moines Office is a field engineering site supporting pilot "
        "customers in Iowa and Nebraska, including GreenField Farms and "
        "AgriNova Cooperative. Employees may also work Remote with manager "
        "approval."
    )

    doc.add_heading("3. Paid Time Off", level=1)
    doc.add_paragraph(
        "Full-time employees accrue 15 days of PTO per year during their "
        "first two years, increasing to 20 days after two years of "
        "service. PTO requests go through Grace Kim in People Operations."
    )

    doc.add_heading("4. Remote Work Policy", level=1)
    doc.add_paragraph(
        "Engineering and Data Science roles may work remotely with VP "
        "approval; Sales and Operations roles are expected on-site at "
        "Austin HQ at least three days per week. Field Applications "
        "Engineers based in Des Moines travel to customer sites as "
        "pilots require, such as the recurring GreenField Farms site "
        "visits."
    )

    doc.add_heading("5. Performance Reviews", level=1)
    doc.add_paragraph(
        "Performance reviews are conducted twice yearly, in January and "
        "July, by each employee's direct manager. Department leads "
        "(Priya Nair for Engineering, Tom Alvarez for Sales) submit "
        "calibrated ratings to Grace Kim for final sign-off by Maria Chen."
    )

    for p in doc.paragraphs:
        for run in p.runs:
            run.font.size = Pt(11)

    doc.save(out_path)


def build_pdf(out_path: Path) -> None:
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    def heading(text: str, size: int) -> None:
        pdf.set_font("Helvetica", "B", size)
        pdf.multi_cell(0, size * 0.7, text, new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    def body(text: str) -> None:
        pdf.set_font("Helvetica", "", 11)
        pdf.multi_cell(0, 7, text, new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    heading("Solstice Robotics -- Q2 2026 Investor Update", 16)
    pdf.ln(2)
    body(
        "Prepared by Maria Chen (CEO) and David Okoye (CTO). Solstice "
        "Robotics closed Q2 2026 with continued progress on both Project "
        "Meridian and Project Halcyon, and is preparing a Series B raise "
        "for early 2027."
    )

    pdf.ln(4)
    heading("Headcount", 13)
    body(
        "Solstice Robotics ended Q2 2026 with 10 employees across Austin "
        "HQ, the Des Moines office, and remote roles, up from 6 a year "
        "earlier. Priya Nair's Engineering team grew to 6 people with the "
        "addition of Ravi Subramaniam as Data Scientist, and Grace Kim's "
        "People Operations team opened a new firmware engineer "
        "requisition to support both flagship projects."
    )

    pdf.ln(4)
    heading("Project Meridian", 13)
    body(
        "The GreenField Farms pilot in Iowa, led on-site by Jordan Blake "
        "with firmware support from Sam Whitfield, remains Solstice's "
        "first paying customer engagement. A firmware power-draw issue "
        "pushed the pilot's second flight window from June 15 to June 29, "
        "but early crop-stress detection results have matched GreenField "
        "Farms' agronomist assessments in 9 of 10 flagged zones."
    )

    pdf.ln(4)
    heading("Project Halcyon", 13)
    body(
        "The swarm-coordination proof of concept with AgriNova "
        "Cooperative in Nebraska wrapped in Q1. Ravi Subramaniam's "
        "updated pathing algorithm reduced flight-path overlap by 18% "
        "versus the naive grid-search baseline used in that proof of "
        "concept. General availability, led by David Okoye, remains "
        "targeted for Q4 2026."
    )

    pdf.ln(4)
    heading("Financing", 13)
    body(
        "Solstice Robotics closed a $12M Series A in January 2023, led by "
        "Terra Ventures with participation from AgFirst Capital. The "
        "company is targeting a Series B in early 2027, timed to the "
        "Project Halcyon general-availability launch, and Tom Alvarez's "
        "Sales team is building a pipeline of reference customers "
        "beyond GreenField Farms and AgriNova Cooperative ahead of that "
        "raise."
    )

    pdf.output(str(out_path))


if __name__ == "__main__":
    import sys

    target_dir = Path(sys.argv[1]) if len(sys.argv) > 1 else OUT_DIR
    build_docx(target_dir / "company_handbook.docx")
    build_pdf(target_dir / "quarterly_report.pdf")
    print(f"Wrote company_handbook.docx and quarterly_report.pdf to {target_dir}")
